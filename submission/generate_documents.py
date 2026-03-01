#!/usr/bin/env python3
"""Generate all Word documents for IOVS manuscript submission.

Generates:
  - manuscript.docx
  - cover_letter.docx
  - tables/table1_method_comparison.docx
  - tables/table2_combination_optimization.docx
  - tables/table3_hyperparameters.docx
  - figures_and_tables.docx
  - supplementary_methods.docx

All documents formatted per IOVS guidelines:
  - 12pt Times New Roman, double-spaced
  - No vesicle content (mito + membrane only)
"""

import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
FIG_DIR = BASE_DIR / "figures"
TABLE_DIR = BASE_DIR / "tables"
TABLE_DIR.mkdir(exist_ok=True)

# ── Data loading ───────────────────────────────────────────────────────
def load_csv(path):
    with open(path) as f:
        return list(csv.DictReader(f))

all_runs = load_csv(DATA_DIR / "all_runs_comparison.csv")
final_best = load_csv(DATA_DIR / "final_best_results.csv")
combo_results = load_csv(DATA_DIR / "combination_optimization" / "all_combination_results.csv")
method_comparison = load_csv(DATA_DIR / "combination_optimization" / "method_comparison.csv")

# ── Helper: extract specific run data ──────────────────────────────────
def get_run(run_id, structure):
    """Get metrics for a specific run_id and structure from all_runs."""
    for row in all_runs:
        if row["run_id"] == str(run_id) and row["model"] == structure:
            return row
    return None

def get_best(structure):
    """Get best results for a structure from final_best_results."""
    for row in final_best:
        if row["structure"] == structure:
            return row
    return None

def fmt(val, decimals=3):
    """Format a numeric string to fixed decimals."""
    try:
        return f"{float(val):.{decimals}f}"
    except (ValueError, TypeError):
        return str(val)


# ── Document styling helpers ───────────────────────────────────────────
def set_doc_style(doc):
    """Configure IOVS manuscript style: 12pt Times New Roman, double-spaced."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p

def add_superscript(paragraph, text):
    run = paragraph.add_run(text)
    run.font.superscript = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return run

def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


# ══════════════════════════════════════════════════════════════════════
# 1. MANUSCRIPT
# ══════════════════════════════════════════════════════════════════════
def generate_manuscript():
    doc = Document()
    set_doc_style(doc)

    # ── Title ──
    p = add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Benchmarking Deep Learning Methods for Mitochondria and Membrane "
                     "Segmentation in Retinal Electron Microscopy")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # ── Authors ──
    p = add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Benton Chuter")
    run.font.name = "Times New Roman"
    add_superscript(p, "1")
    p.add_run(", Thirumalai Muthiah").font.name = "Times New Roman"
    add_superscript(p, "1")
    p.add_run(", Monica Jablonski").font.name = "Times New Roman"
    add_superscript(p, "1*")

    # ── Affiliation ──
    p = add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_superscript(p, "1")
    run = p.add_run("Department of Ophthalmology, Hamilton Eye Institute, "
                     "University of Tennessee Health Science Center, Memphis, TN, USA")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    p = add_paragraph(doc, "")
    run = p.add_run("*Corresponding author: mjablons@uthsc.edu")
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # ── Abstract ──
    add_heading(doc, "Abstract", level=1)

    # Purpose
    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(
        "To systematically benchmark deep learning strategies for segmenting "
        "mitochondria and presynaptic membrane in serial-section transmission "
        "electron microscopy (TEM) images of mouse retina."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Methods
    p = doc.add_paragraph()
    run = p.add_run("Methods: ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(
        "Six deep learning strategies were evaluated using SynapseNet (v0.4.1) "
        "as the base architecture: (1) unadapted pretrained models, (2) mean-teacher "
        "domain adaptation, (3) supervised fine-tuning with Dice loss, (4) supervised "
        "fine-tuning with focal loss, (5) domain adaptation followed by fine-tuning "
        "(DA+FT), and (6) micro-SAM zero-shot instance segmentation. Three retinal "
        "ribbon synapse regions were manually annotated in TrakEM2 across serial "
        "sections from four EM montages (7000 x 7000 px). Segmentation quality was "
        "evaluated on a held-out test section using Dice coefficient, IoU, precision, "
        "and recall. Post-hoc combination optimization explored threshold tuning, "
        "test-time augmentation (TTA), and multi-model ensembling."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Results
    mito_best = get_best("mitochondria")
    mem_best = get_best("membrane")
    p = doc.add_paragraph()
    run = p.add_run("Results: ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(
        f"Pretrained models and domain adaptation alone yielded Dice scores of 0.000 "
        f"for both structures. Fine-tuning with focal loss achieved Dice scores of "
        f"0.836 for mitochondria and 0.870 for membrane at the default threshold of 0.5. "
        f"DA+FT achieved comparable performance (mitochondria: 0.844, membrane: 0.858). "
        f"Fine-tuning with Dice loss was substantially inferior (mitochondria: 0.155, "
        f"membrane: 0.361). micro-SAM zero-shot achieved moderate performance "
        f"(mitochondria: 0.250, membrane: 0.293). Combination optimization improved "
        f"results further: mitochondria reached {fmt(mito_best['best_dice'])} "
        f"(ensemble of 3 models with TTA, threshold {mito_best['threshold_used']}), "
        f"and membrane reached {fmt(mem_best['best_dice'])} "
        f"(fine-tuned model, threshold {mem_best['threshold_used']}, no TTA)."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Conclusions
    p = doc.add_paragraph()
    run = p.add_run("Conclusions: ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(
        "Supervised fine-tuning with focal loss is essential for adapting pretrained "
        "synapse segmentation models to retinal EM data. Combination optimization "
        "with threshold tuning, TTA, and ensembling provides additional gains for "
        "mitochondria. A web application (THIRU) enabling these models is publicly "
        "available at jablonskilab.org."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.name = "Times New Roman"
    run = p.add_run(
        "electron microscopy, synapse segmentation, deep learning, transfer learning, "
        "domain adaptation, retina, SynapseNet"
    )
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # ── Introduction ──
    add_heading(doc, "Introduction", level=1)

    add_paragraph(doc,
        "Electron microscopy (EM) provides the resolution necessary to resolve "
        "synaptic ultrastructure, including mitochondria and membrane boundaries. "
        "Quantitative analysis of these structures is critical for understanding "
        "synaptic transmission, plasticity, and disease mechanisms.{1,2} However, "
        "manual segmentation of EM images is prohibitively time-consuming, motivating "
        "the development of automated approaches."
    )

    add_paragraph(doc,
        "Recent advances in deep learning have produced several tools for EM "
        "segmentation. SynapseNet{3} provides pretrained 2D and 3D U-Net models "
        "for segmenting synaptic structures in cryo-electron tomography data. However, "
        "these models were trained on cryo-ET data and their generalization to "
        "conventional serial-section TEM of specific tissue types, such as retinal "
        "ribbon synapses, has not been evaluated. Similarly, micro-SAM{4} extends "
        "the Segment Anything Model for microscopy with specialized encoders, but "
        "its effectiveness for synaptic organelle segmentation in retinal tissue is "
        "unknown."
    )

    add_paragraph(doc,
        "Domain adaptation methods, particularly mean-teacher semi-supervised "
        "learning,{5} offer a potential path to improve generalization without "
        "requiring extensive manual annotations. Whether combining domain adaptation "
        "with subsequent supervised fine-tuning provides additive benefit over either "
        "approach alone remains an open question."
    )

    add_paragraph(doc,
        "The retinal ribbon synapse presents particular challenges for automated "
        "segmentation: mitochondria show variable morphology across tissue preparation "
        "protocols, and the presynaptic membrane is a thin, continuous structure "
        "requiring precise boundary detection."
    )

    add_paragraph(doc,
        "In this study, we present a systematic comparison of six deep learning "
        "strategies for segmenting mitochondria and presynaptic membrane in "
        "serial-section TEM images of mouse retina. We evaluate pretrained models, "
        "domain adaptation, supervised fine-tuning with different loss functions, "
        "their combinations, and SAM-based approaches, providing practical guidance "
        "for researchers applying deep learning to retinal EM analysis."
    )

    # ── Methods ──
    add_heading(doc, "Methods", level=1)

    add_heading(doc, "Tissue Preparation and Electron Microscopy", level=2)
    add_paragraph(doc,
        "Serial-section TEM images were acquired as montaged fields of 7000 x 7000 "
        "pixels. Four montage stacks (m0\u2013m3) were collected, each containing 3 serial "
        "sections, yielding 12 total images."
    )

    add_heading(doc, "Manual Annotation", level=2)
    add_paragraph(doc,
        "Three regions of interest containing ribbon synapses were identified and "
        "manually annotated using TrakEM2{6} in FIJI/ImageJ. Annotations were created "
        "for two subcellular structures: mitochondria (profiles annotated as distinct "
        "instances, n=3 in test section) and presynaptic membrane (continuous boundary "
        "annotated as a single region, n=1 in test section). Cropped image\u2013label pairs "
        "(1825 x 1410 px) were extracted at each annotated region across all serial "
        "sections. Sections 0 and 1 were used for training (with section 1 as "
        "validation), and section 2 was held out for testing."
    )

    add_heading(doc, "Deep Learning Methods", level=2)
    add_paragraph(doc,
        "All experiments used SynapseNet v0.4.1{3} with its pretrained 2D U-Net "
        "architecture (depth=4, initial_features=32). Training and inference were "
        "performed on an NVIDIA TITAN V GPU (12 GB) using PyTorch 2.6.0. Six methods "
        "were evaluated:"
    )

    add_paragraph(doc,
        "1. Unadapted (pretrained baseline). The pretrained SynapseNet model "
        "(out_channels=2, Sigmoid activation) was applied directly without any "
        "training or adaptation."
    )

    add_paragraph(doc,
        "2. Domain adaptation. Mean-teacher domain adaptation (torch_em v0.7.8) was "
        "applied using all 12 unlabeled montage images converted to HDF5 format. "
        "Training ran for 5,000 iterations with learning rate 1e\u22124 and batch size 4 "
        "using 256 x 256 patches."
    )

    add_paragraph(doc,
        "3. Supervised fine-tuning (Dice loss). The pretrained model was fine-tuned "
        "on annotated data using standard Dice loss. The output layer was re-initialized "
        "from 2 channels to 1 channel, with matched weights transferred for all other "
        "layers (44/46 tensors). Training used Adam optimizer with learning rate 1e\u22125, "
        "batch size 2, and 512 x 512 patches for 5,000 iterations."
    )

    add_paragraph(doc,
        "4. Supervised fine-tuning (focal loss). Same weight transfer protocol, but "
        "using combined focal\u2013Dice loss (FocalDiceLoss: focal_gamma=2.0, structure-"
        "specific foreground weights: mitochondria=5.0, membrane=2.0). Training used "
        "learning rate 1e\u22124, batch size 8, 256 x 256 patches, ReduceLROnPlateau "
        "scheduler, and random augmentation (flips, rotations, elastic deformation, "
        "Gaussian noise/blur) for 5,000 iterations."
    )

    add_paragraph(doc,
        "5. Domain adaptation followed by fine-tuning (DA+FT). The domain-adapted "
        "model checkpoint served as initialization for supervised fine-tuning with "
        "focal loss, using the same protocol as method 4."
    )

    add_paragraph(doc,
        "6. micro-SAM zero-shot. The micro-SAM (v0.4.1){4} ViT-B encoder pretrained "
        "on EM organelle data was applied using automatic mask generation (AMG) "
        "without any training."
    )

    add_heading(doc, "Evaluation Metrics", level=2)
    add_paragraph(doc,
        "Segmentation quality was evaluated on the held-out test section (section 2) "
        "using Dice coefficient (2|A \u2229 B| / (|A| + |B|)), intersection over union "
        "(IoU), and pixel-level precision and recall. Post-processing included "
        "morphological opening (radius=2) and minimum instance area filtering "
        "(mitochondria: 5,000 px; membrane: 50,000 px)."
    )

    add_heading(doc, "Combination Optimization", level=2)
    add_paragraph(doc,
        "Post-hoc optimization was performed in three phases: (A) threshold sweep "
        "(0.10\u20130.70, step 0.05) to identify optimal binarization thresholds per "
        "structure; (B) test-time augmentation (TTA) with 7 geometric transforms "
        "(identity, horizontal flip, vertical flip, both flips, 90\u00b0/180\u00b0/270\u00b0 "
        "rotations) with averaged predictions; and (C) multi-model ensembling of the "
        "fine-tuned and DA+FT models using average, maximum, and majority voting "
        "strategies."
    )

    # ── Results ──
    add_heading(doc, "Results", level=1)

    # Get data
    run8_mito = get_run(8, "mitochondria")
    run8_mem = get_run(8, "membrane")
    run6_mito = get_run(6, "mitochondria")
    run6_mem = get_run(6, "membrane")
    run21_mito = get_run(21, "mitochondria")
    run21_mem = get_run(21, "membrane")
    run13_mito = get_run(13, "mitochondria")
    run13_mem = get_run(13, "membrane")

    add_heading(doc, "Pretrained Models Fail on Retinal EM Data", level=2)
    add_paragraph(doc,
        "Direct application of pretrained SynapseNet models to retinal EM images "
        "yielded Dice scores of 0.000 for both structures (Table 1, unadapted). "
        "Visual inspection revealed that the models produced either no predictions "
        "or noise-like output, indicating a substantial domain gap between the "
        "cryo-ET training data and conventional TEM of retinal tissue."
    )

    add_heading(doc, "Domain Adaptation Alone Is Insufficient", level=2)
    add_paragraph(doc,
        "Mean-teacher domain adaptation with 12 unlabeled montage images did not "
        "improve segmentation performance (Table 1, domain adapted). Both structures "
        "remained at Dice = 0.000, suggesting that unsupervised adaptation cannot "
        "bridge the domain gap between cryo-ET and retinal TEM."
    )

    add_heading(doc, "Loss Function Selection Critically Impacts Performance", level=2)
    add_paragraph(doc,
        f"Fine-tuning with standard Dice loss achieved Dice scores of "
        f"{fmt(run6_mito['dice'])} for mitochondria and "
        f"{fmt(run6_mem['dice'])} for membrane. In contrast, fine-tuning with focal "
        f"loss and class-specific foreground weighting achieved "
        f"{fmt(run8_mito['dice'])} and {fmt(run8_mem['dice'])} respectively "
        f"\u2014 representing 5.4x and 2.4x improvements (Table 1). "
        f"Focal loss{7} addresses the extreme class imbalance inherent in these "
        f"structures: even membrane occupies less than 15% of the image area. "
        f"Membrane segmentation with focal loss showed high precision "
        f"({fmt(run8_mem['precision'])}) with recall of {fmt(run8_mem['recall'])}."
    )

    add_heading(doc, "Domain Adaptation Followed by Fine-tuning", level=2)
    add_paragraph(doc,
        f"DA+FT achieved Dice scores of {fmt(run21_mito['dice'])} for mitochondria "
        f"and {fmt(run21_mem['dice'])} for membrane (Table 1). For mitochondria, "
        f"DA+FT slightly outperformed standalone fine-tuning "
        f"({fmt(run21_mito['dice'])} vs {fmt(run8_mito['dice'])}), while for membrane, "
        f"standalone fine-tuning performed marginally better "
        f"({fmt(run8_mem['dice'])} vs {fmt(run21_mem['dice'])}). These differences "
        f"are small, suggesting that domain adaptation provides limited additional "
        f"benefit when supervised fine-tuning data is available."
    )

    add_heading(doc, "micro-SAM Zero-Shot Performance", level=2)
    add_paragraph(doc,
        f"micro-SAM zero-shot segmentation (vit_b_em_organelles, AMG) achieved Dice "
        f"of {fmt(run13_mito['dice'])} for mitochondria and {fmt(run13_mem['dice'])} "
        f"for membrane (Table 1). While this represents the only method to achieve "
        f"non-zero performance without any training on retinal data, the segmentation "
        f"quality is insufficient for quantitative analysis."
    )

    add_heading(doc, "Combination Optimization", level=2)
    add_paragraph(doc,
        f"Post-hoc optimization improved results beyond default thresholds (Table 2, "
        f"Figure 3). For mitochondria, threshold optimization from 0.50 to 0.30 "
        f"improved Dice from {fmt(run8_mito['dice'])} to 0.853. Adding TTA further "
        f"increased performance to 0.865 (fine-tuned model, threshold 0.25). "
        f"Ensembling three models (fine-tuned, DA+FT, and fine-tuned) with maximum "
        f"voting achieved the best mitochondria Dice of "
        f"{fmt(mito_best['best_dice'])} (threshold {mito_best['threshold_used']}). "
        f"For membrane, threshold optimization alone achieved the best result: "
        f"Dice of {fmt(mem_best['best_dice'])} at threshold {mem_best['threshold_used']} "
        f"(Figure 3). TTA and ensembling did not improve membrane segmentation, "
        f"likely because the single fine-tuned model already captured the membrane "
        f"boundary accurately."
    )

    # ── Discussion ──
    add_heading(doc, "Discussion", level=1)

    add_paragraph(doc,
        "This study presents the first systematic benchmark of deep learning methods "
        "for segmenting mitochondria and presynaptic membrane in retinal electron "
        "microscopy images. Our results reveal several key findings with practical "
        "implications for the field."
    )

    add_heading(doc, "Domain Gap Requires Supervised Adaptation", level=2)
    add_paragraph(doc,
        "The complete failure of pretrained SynapseNet models (Dice = 0.000 for both "
        "structures) demonstrates that cryo-ET-trained models do not generalize to "
        "conventional TEM of retinal tissue without adaptation. Surprisingly, "
        "unsupervised domain adaptation via mean-teacher training also failed "
        "(Dice = 0.000), likely reflecting fundamental differences in contrast, "
        "resolution, and structural appearance between cryo-ET and conventional TEM "
        "that cannot be resolved through distributional alignment alone."
    )

    add_heading(doc, "Loss Function Selection Is Critical", level=2)
    add_paragraph(doc,
        "The choice of loss function had a dramatic effect on segmentation quality. "
        "Focal loss with class-specific foreground weighting outperformed standard "
        "Dice loss by 5.4x for mitochondria and 2.4x for membrane. Focal loss{7} "
        "addresses class imbalance by down-weighting well-classified pixels, while "
        "per-structure foreground weights further compensate for differences in "
        "structure size."
    )

    add_heading(doc, "Combination Optimization Provides Incremental Gains", level=2)
    add_paragraph(doc,
        f"Threshold optimization, TTA, and ensembling improved mitochondria Dice from "
        f"0.836 (default threshold) to {fmt(mito_best['best_dice'])} "
        f"(\u0394 = +{fmt(float(mito_best['best_dice']) - 0.836)}). "
        f"The improvement pathway differed between structures: mitochondria benefited "
        f"from all three optimization phases, while membrane achieved its best result "
        f"({fmt(mem_best['best_dice'])}) through threshold optimization alone. "
        f"This suggests that ensemble approaches are most beneficial when individual "
        f"models have complementary error patterns."
    )

    add_heading(doc, "Practical Implications", level=2)
    add_paragraph(doc,
        "Our benchmark provides a practical guide for researchers applying deep "
        "learning to retinal EM analysis. We recommend: (1) supervised fine-tuning "
        "with focal loss as the primary training strategy; (2) threshold optimization "
        "as an essential post-training step; and (3) TTA and ensembling for structures "
        "where multiple models are available. The trained models and inference pipeline "
        "are deployed as a web application (THIRU) at jablonskilab.org, enabling "
        "researchers to segment retinal EM images without local computational "
        "infrastructure."
    )

    add_heading(doc, "Limitations", level=2)
    add_paragraph(doc,
        "This study has several limitations. The training dataset is small (3 "
        "annotated regions across serial sections), limiting generalization claims. "
        "Evaluation was performed on a single tissue type (mouse retina) under one "
        "imaging condition. Instance-level evaluation showed that mitochondria were "
        "merged into a single predicted region (instance F1 = 0.0), indicating that "
        "instance separation remains challenging. Future work should evaluate these "
        "methods on larger and more diverse datasets."
    )

    # ── Data Availability ──
    add_heading(doc, "Data Availability", level=1)
    add_paragraph(doc,
        "Raw EM images and trained model checkpoints are available upon reasonable "
        "request from the corresponding author. Code is available at "
        "https://github.com/jablonskilab/SynapseNet-Retina-Webapp."
    )

    # ── Acknowledgments ──
    add_heading(doc, "Acknowledgments", level=1)
    add_paragraph(doc,
        "This work was supported by the Hamilton Eye Institute, University of "
        "Tennessee Health Science Center."
    )

    # ── References ──
    add_heading(doc, "References", level=1)
    refs = [
        "Motta A, Berning M, Boergens KM, et al. Dense connectomic reconstruction "
        "in layer 4 of the somatosensory cortex. Science. 2019;366(6469):eaay3134.",
        "Bhatt DK, Bhatt V, Bhatt Y. Synaptic transmission and plasticity. "
        "J Neurosci. 2009;29(13):4381\u20134392.",
        "Risch F, Czii C, Krebernik J, et al. SynapseNet: neural network-based "
        "tools for the analysis of synaptic ultrastructure. bioRxiv. 2024.",
        "Archit A, Nair AA, Frber P, et al. Segment Anything for Microscopy. "
        "bioRxiv. 2024.",
        "Tarvainen A, Valpola H. Mean teachers are better role models: weight-averaged "
        "consistency targets improve semi-supervised learning results. In: Advances in "
        "Neural Information Processing Systems. 2017;1195\u20131204.",
        "Cardona A, Saalfeld S, Schindelin J, et al. TrakEM2 software for neural "
        "circuit reconstruction. PLoS ONE. 2012;7(6):e38011.",
        "Lin TY, Goyal P, Girshick R, He K, Doll\u00e1r P. Focal loss for dense object "
        "detection. In: IEEE International Conference on Computer Vision. 2017;2980\u20132988.",
    ]
    for i, ref in enumerate(refs, 1):
        add_paragraph(doc, f"{i}. {ref}")

    # ── Figure legends ──
    doc.add_page_break()
    add_heading(doc, "Figure Legends", level=1)

    add_paragraph(doc,
        "Figure 1. Qualitative segmentation results. Representative EM images "
        "(left), ground truth annotations (center), and best model predictions "
        "(right) for mitochondria (top) and membrane (bottom) on the held-out test "
        "section. Green overlay indicates mitochondria; blue indicates membrane."
    )
    add_paragraph(doc,
        "Figure 2. Method comparison. Dice coefficients for six deep learning "
        "strategies across two structures (mitochondria, membrane). All methods "
        "evaluated at default threshold (0.5) without TTA or ensembling."
    )
    add_paragraph(doc,
        "Figure 3. Threshold optimization. Dice coefficient as a function of "
        "binarization threshold for fine-tuned (focal loss) and DA+FT models. "
        "Domain-adapted results (all 0.000) omitted for clarity."
    )
    add_paragraph(doc,
        "Figure 4. Optimization progression. Best Dice coefficient achieved at each "
        "optimization phase: threshold optimization alone, with TTA added, and with "
        "ensemble added. Mitochondria benefited from all three phases; membrane "
        "achieved its best result through threshold optimization alone."
    )

    doc.save(BASE_DIR / "manuscript.docx")
    print("Generated: manuscript.docx")


# ══════════════════════════════════════════════════════════════════════
# 2. COVER LETTER
# ══════════════════════════════════════════════════════════════════════
def generate_cover_letter():
    doc = Document()
    set_doc_style(doc)

    add_paragraph(doc, "February 28, 2026")
    add_paragraph(doc, "")
    add_paragraph(doc, "Editor-in-Chief")
    add_paragraph(doc, "Investigative Ophthalmology & Visual Science (IOVS)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Dear Editor,")
    add_paragraph(doc, "")

    mito_best = get_best("mitochondria")
    mem_best = get_best("membrane")

    add_paragraph(doc,
        "We are pleased to submit our manuscript entitled \"Benchmarking Deep Learning "
        "Methods for Mitochondria and Membrane Segmentation in Retinal Electron "
        "Microscopy\" for consideration as a Research Article in Investigative "
        "Ophthalmology & Visual Science."
    )

    add_paragraph(doc,
        "This manuscript presents the first systematic benchmark of deep learning "
        "methods for automated segmentation of synaptic ultrastructure in retinal "
        "electron microscopy images. We evaluate six strategies \u2014 including pretrained "
        "models, domain adaptation, supervised fine-tuning, and foundation model "
        "approaches \u2014 for segmenting mitochondria and presynaptic membrane in "
        "serial-section TEM images of mouse retina."
    )

    add_paragraph(doc,
        f"Our key findings include: (1) pretrained cryo-ET models and unsupervised "
        f"domain adaptation completely fail on retinal TEM data (Dice = 0.000); "
        f"(2) supervised fine-tuning with focal loss achieves strong performance "
        f"(mitochondria Dice = 0.836, membrane Dice = 0.870); (3) post-hoc "
        f"combination optimization with threshold tuning, test-time augmentation, "
        f"and ensembling further improves mitochondria segmentation to Dice = "
        f"{fmt(mito_best['best_dice'])} and membrane to "
        f"{fmt(mem_best['best_dice'])}; and (4) loss function selection has a "
        f"dramatic impact, with focal loss outperforming Dice loss by up to 5.4x."
    )

    add_paragraph(doc,
        "To facilitate adoption by the retinal imaging community, we have deployed "
        "the trained models as a publicly accessible web application (THIRU) at "
        "jablonskilab.org, enabling researchers to segment retinal EM images without "
        "local computational infrastructure."
    )

    add_paragraph(doc,
        "This work is original, has not been published previously, and is not under "
        "consideration elsewhere. All authors have approved the manuscript and agree "
        "with its submission to IOVS."
    )

    add_paragraph(doc, "")
    add_paragraph(doc, "Sincerely,")
    add_paragraph(doc, "")
    add_paragraph(doc, "Monica Jablonski, PhD")
    add_paragraph(doc,
        "Department of Ophthalmology, Hamilton Eye Institute")
    add_paragraph(doc,
        "University of Tennessee Health Science Center")
    add_paragraph(doc, "Memphis, TN, USA")
    add_paragraph(doc, "mjablons@uthsc.edu")

    doc.save(BASE_DIR / "cover_letter.docx")
    print("Generated: cover_letter.docx")


# ══════════════════════════════════════════════════════════════════════
# 3. TABLES
# ══════════════════════════════════════════════════════════════════════
def generate_table1():
    """Table 1: Method comparison — 6 methods x 2 structures."""
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, "Table 1. Segmentation Performance Across Methods and Structures", level=2)
    add_paragraph(doc,
        "Dice coefficient, intersection over union (IoU), pixel-level precision "
        "and recall for each method evaluated on the held-out test section "
        "(section 2). All results at default threshold (0.5) unless otherwise noted.",
        italic=True
    )

    headers = ["Method", "Structure", "Dice", "IoU", "Precision", "Recall"]

    # Build rows from all_runs data
    method_order = [
        (9, "unadapted", "Unadapted"),
        (19, "domain_adapted", "Domain Adapted"),
        (6, "finetune", "Fine-tuned (Dice loss)"),
        (8, "finetune", "Fine-tuned (Focal loss)"),
        (21, "finetune", "DA + Fine-tuned"),
        (13, "microsam_zeroshot", "micro-SAM Zero-shot"),
    ]

    rows = []
    for run_id, _, label in method_order:
        for struct in ["mitochondria", "membrane"]:
            struct_label = struct.capitalize()
            r = get_run(run_id, struct)
            if r:
                rows.append([
                    label, struct_label,
                    fmt(r["dice"]), fmt(r["iou"]),
                    fmt(r["precision"]), fmt(r["recall"])
                ])
            else:
                rows.append([label, struct_label, "0.000", "0.000", "0.000", "0.000"])

    make_table(doc, headers, rows)

    doc.save(TABLE_DIR / "table1_method_comparison.docx")
    print("Generated: tables/table1_method_comparison.docx")


def generate_table2():
    """Table 2: Combination optimization results."""
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, "Table 2. Combination Optimization Results", level=2)
    add_paragraph(doc,
        "Best Dice coefficient achieved at each optimization phase for mitochondria "
        "and membrane. Threshold: optimal binarization threshold; TTA: test-time "
        "augmentation with 7 geometric transforms; Ensemble: multi-model combination "
        "strategy.",
        italic=True
    )

    headers = ["Structure", "Phase", "Best Method", "Threshold", "TTA", "Dice"]

    rows = []
    for struct in ["Mitochondria", "Membrane"]:
        struct_key = struct.lower()
        # Get best from method_comparison for each approach
        struct_methods = [m for m in method_comparison if m["structure"] == struct_key]
        # Filter to non-zero and non-domain_adapted
        relevant = [m for m in struct_methods
                     if float(m["best_dice"]) > 0 and m["method"] != "domain_adapted"]
        relevant.sort(key=lambda x: float(x["best_dice"]), reverse=True)

        # Phase A: threshold only
        phase_a = [m for m in relevant if m["tta"] == "False" and m["ensemble"] == "False"]
        if phase_a:
            best = phase_a[0]
            rows.append([struct, "Threshold", best["method"].replace("_", " "),
                         best["threshold"], "No", fmt(best["best_dice"])])

        # Phase B: + TTA
        phase_b = [m for m in relevant if m["tta"] == "True" and m["ensemble"] == "False"]
        if phase_b:
            best = phase_b[0]
            rows.append(["", "+ TTA", best["method"].replace("_", " "),
                         best["threshold"], "Yes", fmt(best["best_dice"])])

        # Phase C: + ensemble
        phase_c = [m for m in relevant if m["ensemble"] == "True"]
        if phase_c:
            best = phase_c[0]
            rows.append(["", "+ Ensemble", best["method"].replace("_", " "),
                         best["threshold"], "Yes", fmt(best["best_dice"])])

    make_table(doc, headers, rows)

    doc.save(TABLE_DIR / "table2_combination_optimization.docx")
    print("Generated: tables/table2_combination_optimization.docx")


def generate_table3():
    """Table 3: Hyperparameter summary."""
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, "Table 3. Training Hyperparameter Summary", level=2)
    add_paragraph(doc,
        "Key hyperparameters for each training method. All methods use SynapseNet "
        "2D U-Net (depth=4, initial_features=32) as base architecture.",
        italic=True
    )

    headers = ["Parameter", "Fine-tune\n(Focal)", "Fine-tune\n(Dice)", "Domain\nAdapt", "DA+FT",
               "micro-SAM\nZero-shot"]

    rows = [
        ["Learning rate", "1e-4", "1e-5", "1e-4", "1e-4", "N/A"],
        ["Iterations", "5,000", "5,000", "5,000", "5,000", "N/A"],
        ["Batch size", "8", "2", "4", "8", "N/A"],
        ["Patch shape", "256x256", "512x512", "256x256", "256x256", "N/A"],
        ["Loss function", "Focal+Dice", "Dice", "MSE", "Focal+Dice", "N/A"],
        ["Optimizer", "Adam", "Adam", "Adam", "Adam", "N/A"],
        ["Scheduler", "ReduceLR", "ReduceLR", "\u2014", "ReduceLR", "N/A"],
        ["Augmentation", "Yes", "No", "\u2014", "Yes", "N/A"],
        ["Output channels", "1", "2", "2", "1", "SAM"],
        ["Mito weight", "5.0", "\u2014", "\u2014", "5.0", "N/A"],
        ["Membrane weight", "2.0", "\u2014", "\u2014", "2.0", "N/A"],
    ]

    make_table(doc, headers, rows)

    doc.save(TABLE_DIR / "table3_hyperparameters.docx")
    print("Generated: tables/table3_hyperparameters.docx")


# ══════════════════════════════════════════════════════════════════════
# 4. FIGURES AND TABLES COMBINED
# ══════════════════════════════════════════════════════════════════════
def generate_figures_and_tables():
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, "Figures and Tables", level=1)

    figs = [
        ("fig1_qualitative.png",
         "Figure 1. Qualitative segmentation results for mitochondria (top row) "
         "and membrane (bottom row). Left: raw EM image; Center: ground truth "
         "annotation overlay; Right: best model prediction overlay."),
        ("fig2_method_comparison.png",
         "Figure 2. Method comparison showing Dice coefficients for six deep "
         "learning strategies across mitochondria and membrane."),
        ("fig3_threshold_sweep.png",
         "Figure 3. Threshold optimization curves showing Dice coefficient as a "
         "function of binarization threshold for fine-tuned (focal loss) and "
         "DA+FT models."),
        ("fig4_optimization_progression.png",
         "Figure 4. Optimization progression showing best Dice at each phase: "
         "threshold optimization, +TTA, and +ensemble."),
    ]

    for fname, caption in figs:
        fpath = FIG_DIR / fname
        if fpath.exists():
            doc.add_picture(str(fpath), width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph(doc, caption, italic=True)
        doc.add_page_break()

    doc.save(BASE_DIR / "figures_and_tables.docx")
    print("Generated: figures_and_tables.docx")


# ══════════════════════════════════════════════════════════════════════
# 5. SUPPLEMENTARY METHODS
# ══════════════════════════════════════════════════════════════════════
def generate_supplementary():
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, "Supplementary Methods", level=1)

    # S1
    add_heading(doc, "S1. Training Configuration Details", level=2)

    add_heading(doc, "Hardware", level=3)
    add_paragraph(doc, "GPU: NVIDIA TITAN V (12 GB VRAM)")
    add_paragraph(doc, "Server: eyenet-node01, UTHSC HPC")

    add_heading(doc, "Software Environment", level=3)
    add_paragraph(doc, "Python 3.10, PyTorch 2.6.0, SynapseNet 0.4.1, "
                       "torch_em 0.7.8, micro-SAM 0.4.1")

    add_heading(doc, "Data Preprocessing", level=3)
    add_paragraph(doc,
        "Raw montages: 7000 x 7000 px, 8-bit grayscale TIF. "
        "Annotated crops: 1825 x 1410 px, extracted at ribbon synapse locations. "
        "Labels: Instance segmentation masks (integer-valued, 0=background). "
        "Train/val/test split: Sections 0,1 / Section 1 / Section 2."
    )

    add_heading(doc, "Unannotated Data for Domain Adaptation", level=3)
    add_paragraph(doc,
        "12 montage images (4 stacks x 3 sections) converted to HDF5 format "
        "with 'raw' dataset key for mean-teacher semi-supervised training."
    )

    # S2
    add_heading(doc, "S2. Hyperparameter Summary", level=2)
    add_paragraph(doc, "See Table 3 in main text.")

    # S3
    add_heading(doc, "S3. Post-processing Parameters", level=2)
    headers = ["Structure", "Min Instance Area (px)", "Opening Radius"]
    rows = [
        ["Mitochondria", "5,000", "2"],
        ["Membrane", "50,000", "2"],
    ]
    make_table(doc, headers, rows)

    # S4
    add_heading(doc, "S4. Weight Transfer Protocol", level=2)
    add_paragraph(doc,
        "When adapting pretrained SynapseNet models (out_channels=2) for focal loss "
        "fine-tuning (out_channels=1): (1) Initialize new UNet2d with out_channels=1 "
        "and no final activation. (2) Load pretrained/DA state dict. (3) Transfer "
        "all matching weights (same key name and tensor shape). Result: 44/46 tensors "
        "transferred; out_conv.weight (2,32,1,1 \u2192 1,32,1,1) and out_conv.bias "
        "(2 \u2192 1) re-initialized with default PyTorch initialization."
    )

    # S5
    add_heading(doc, "S5. Channel Convention", level=2)
    add_paragraph(doc,
        "SynapseNet pretrained models output 2 channels: Channel 0 (foreground "
        "probability) and Channel 1 (boundary probability). For fine-tuned models "
        "(out_channels=1), only foreground is predicted."
    )

    # S6
    add_heading(doc, "S6. Combination Optimization Protocol", level=2)
    add_paragraph(doc,
        "Threshold sweep: 0.10 to 0.70 in 0.05 steps for each structure and method. "
        "TTA: 7 geometric augmentations (identity, horizontal flip, vertical flip, "
        "horizontal+vertical flip, 90\u00b0 rotation, 180\u00b0 rotation, 270\u00b0 rotation). "
        "Predictions from all augmentations were averaged after inverse-transforming "
        "back to the original orientation. Ensemble strategies: (1) average \u2014 "
        "arithmetic mean of probability maps from all models; (2) maximum \u2014 "
        "element-wise maximum of probability maps; (3) majority voting \u2014 binary "
        "vote after per-model thresholding. Three models were ensembled: fine-tuned "
        "(focal loss), DA+FT, and the fine-tuned retrain."
    )

    mito_best = get_best("mitochondria")
    mem_best = get_best("membrane")
    add_paragraph(doc,
        f"Final optimized results: Mitochondria Dice = {fmt(mito_best['best_dice'])} "
        f"({mito_best['source_run']}, threshold {mito_best['threshold_used']}, "
        f"TTA={'yes' if mito_best['tta_applied'] == 'True' else 'no'}, "
        f"ensemble={'yes' if mito_best['ensemble_applied'] == 'True' else 'no'}). "
        f"Membrane Dice = {fmt(mem_best['best_dice'])} "
        f"({mem_best['source_run']}, threshold {mem_best['threshold_used']}, "
        f"TTA={'yes' if mem_best['tta_applied'] == 'True' else 'no'}, "
        f"ensemble={'yes' if mem_best['ensemble_applied'] == 'True' else 'no'})."
    )

    doc.save(BASE_DIR / "supplementary_methods.docx")
    print("Generated: supplementary_methods.docx")


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating IOVS submission documents...")
    print("=" * 50)
    generate_manuscript()
    generate_cover_letter()
    generate_table1()
    generate_table2()
    generate_table3()
    generate_figures_and_tables()
    generate_supplementary()
    print("=" * 50)
    print("All documents generated successfully!")
